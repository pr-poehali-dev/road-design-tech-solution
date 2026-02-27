"""Генерация КП через ИИ — асинхронная очередь через БД, поллинг статуса"""
import json
import os
import base64
import io
import httpx
import psycopg2
import threading


def get_db():
    return psycopg2.connect(os.environ['DATABASE_URL'])


def extract_text_from_pdf(data: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        texts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
        return '\n'.join(texts)
    except Exception as e:
        return f'[Ошибка чтения PDF: {e}]'


def extract_text_from_docx(data: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        parts = []

        def get_cell_text(cell):
            return ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())

        for block in doc.element.body:
            tag = block.tag.split('}')[-1] if '}' in block.tag else block.tag
            if tag == 'p':
                para = None
                for p in doc.paragraphs:
                    if p._element is block:
                        para = p
                        break
                if para and para.text.strip():
                    style = para.style.name if para.style else ''
                    if 'Heading' in style or 'heading' in style:
                        parts.append(f'\n## {para.text.strip()}')
                    else:
                        parts.append(para.text.strip())
            elif tag == 'tbl':
                for tbl in doc.tables:
                    if tbl._element is block:
                        rows = tbl.rows
                        if not rows:
                            continue
                        header_cells = [get_cell_text(c) for c in rows[0].cells]
                        has_header = any(header_cells)
                        if has_header:
                            parts.append('\nТАБЛИЦА:')
                            parts.append(' | '.join(header_cells))
                            parts.append('-' * 40)
                        for row in (rows[1:] if has_header else rows):
                            row_texts = [get_cell_text(c) for c in row.cells]
                            non_empty = [t for t in row_texts if t]
                            if non_empty:
                                if has_header:
                                    pairs = []
                                    for h, v in zip(header_cells, row_texts):
                                        if v:
                                            pairs.append(f'{h}: {v}' if h else v)
                                    parts.append(', '.join(pairs))
                                else:
                                    parts.append(' | '.join(non_empty))
                        break

        return '\n'.join(parts)
    except Exception as e:
        return f'[Ошибка чтения DOCX: {e}]'


PRICE_LIST = [
    {"code": "ОП-ОКС.БЗУ", "name": "Благоустройство земельного участка (БЗУ под ключ)", "unit": "га", "price_per_unit": 500000, "min_order_sum": 500000, "special_rules": "Округление площади: 1.1–1.4 как 1 га; 1.5–1.9 как 2 га"},
    {"code": "ОП-ППТ понижающие условия", "name": "ППТ, гос, понижающие условия", "unit": "га", "price_per_unit": 215000, "special_rules": "Коэффициент 0.15 для понижающих условий"},
    {"code": "ОП-ППТ госконтракт", "name": "ППТ, госконтракт", "unit": "га", "price_per_unit": 450000, "special_rules": "Коэффициент 0.3 для госконтрактов"},
    {"code": "ОП-ППТ коммерческий", "name": "ППТ, коммерческий контракт", "unit": "га", "price_per_unit": 1500000, "special_rules": "Коэффициент 1 для коммерческих контрактов"},
    {"code": "ОП-ПМТ понижающие условия", "name": "ПМТ, гос, понижающие условия", "unit": "га", "price_per_unit": 215000},
    {"code": "ОП-ПМТ госконтракт", "name": "ПМТ, госконтракт", "unit": "га", "price_per_unit": 450000},
    {"code": "ОП-ПМТ коммерческий", "name": "ПМТ, коммерческий контракт", "unit": "га", "price_per_unit": 1500000},
    {"code": "ИИ-ИГДИ", "name": "Инженерно-геодезические изыскания", "unit": "га", "price_per_unit": 40000},
    {"code": "ИИ-ИЭИ", "name": "Инженерно-экологические изыскания", "unit": "га", "price_per_unit": 65000},
    {"code": "ИИ-ОБМ", "name": "Обследование строительных конструкций (обмерные работы)", "unit": "м³", "price_per_unit": 500},
    {"code": "ОП-ПГО до 100", "name": "ПГО до 100 га", "unit": "га", "price_per_unit": 24000},
    {"code": "ОП-ПГО 100-500", "name": "ПГО 100–500 га", "unit": "га", "price_per_unit": 23000},
    {"code": "ОП-ПГО 500-2000", "name": "ПГО 500–2000 га", "unit": "га", "price_per_unit": 22000},
    {"code": "ОП-ПГО 2000-5000", "name": "ПГО 2000–5000 га", "unit": "га", "price_per_unit": 21000},
    {"code": "ОП-ПГО более 5000", "name": "ПГО более 5000 га", "unit": "га", "price_per_unit": 20000},
    {"code": "ОП-ЛИК.КОНС.ГДП", "name": "Ликвидация/консервация/рекультивация горнодобывающих участков", "unit": "га", "price_per_unit": 100000},
    {"code": "СОПГЭ", "name": "Сопровождение экспертизы (ГЭ/НГЭ)", "unit": "комплект", "price_per_unit": 200000},
    {"code": "ОП-ЛО.НС.АД до 1000", "name": "Дорога НС, до 1000 м²", "unit": "м²", "price_per_unit": 700},
    {"code": "ОП-ЛО.НС.АД 1000-5000", "name": "Дорога НС, 1000–5000 м²", "unit": "м²", "price_per_unit": 600},
    {"code": "ОП-ЛО.НС.АД 5000-10000", "name": "Дорога НС, 5000–10000 м²", "unit": "м²", "price_per_unit": 550},
    {"code": "ОП-ЛО.НС.АД более 10000", "name": "Дорога НС, более 10000 м²", "unit": "м²", "price_per_unit": 500},
    {"code": "ОП-ЛО.РЕК.АД до 1000", "name": "Дорога реконструкция, до 1000 м²", "unit": "м²", "price_per_unit": 2000},
    {"code": "ОП-ЛО.РЕК.АД 1000-5000", "name": "Дорога реконструкция, 1000–5000 м²", "unit": "м²", "price_per_unit": 1400},
    {"code": "ОП-ОКС.НС.МОСТ", "name": "Мостовое сооружение, новое", "unit": "м²", "price_per_unit": 11000},
    {"code": "ОП-ОКС.РЕК.МОСТ", "name": "Мостовое сооружение, реконструкция", "unit": "м²", "price_per_unit": 9000},
    {"code": "ОП-ДП до 100", "name": "Дизайн-проект интерьеров до 100 м²", "unit": "м²", "price_per_unit": 5000},
    {"code": "ОП-ДП 100-200", "name": "Дизайн-проект 100–200 м²", "unit": "м²", "price_per_unit": 4000},
    {"code": "ОП-ДП 200-500", "name": "Дизайн-проект 200–500 м²", "unit": "м²", "price_per_unit": 3000},
    {"code": "ОП-ОКС.НС до 100", "name": "ОКС — новое строительство, до 100 м³", "unit": "м³", "price_per_unit": 5000},
    {"code": "ОП-ОКС.НС 100-500", "name": "ОКС — новое строительство, 100–500 м³", "unit": "м³", "price_per_unit": 4000},
    {"code": "ОП-ОКС.НС 500-5000", "name": "ОКС — новое строительство, 500–5000 м³", "unit": "м³", "price_per_unit": 2000},
    {"code": "ОП-ОКС.НС более 5000", "name": "ОКС — новое строительство, более 5000 м³", "unit": "м³", "price_per_unit": 1500},
    {"code": "ОП-ОКС.РЕК до 100", "name": "ОКС — реконструкция, до 100 м³", "unit": "м³", "price_per_unit": 3000},
    {"code": "ОП-ОКС.РЕК 100-500", "name": "ОКС — реконструкция, 100–500 м³", "unit": "м³", "price_per_unit": 2500},
    {"code": "ОП-ОКС.КР до 100", "name": "ОКС — капитальный ремонт, до 100 м³", "unit": "м³", "price_per_unit": 2000},
    {"code": "ОП-ОКС.КР 100-500", "name": "ОКС — капитальный ремонт, 100–500 м³", "unit": "м³", "price_per_unit": 1500},
    {"code": "ОП-АН", "name": "Авторский надзор", "unit": "месяц", "price_per_unit": 100000},
    {"code": "ОП-НТС", "name": "Научно-техническое сопровождение (НТС)", "unit": "месяц", "price_per_unit": 1500000, "min_order_sum": 1500000},
    {"code": "ОП-ТНЗ 1 сотрудник", "name": "Технадзор/техсопровождение, 1 сотрудник", "unit": "чел.-месяц", "price_per_unit": 350000},
    {"code": "ОП-ТНЗ 2-5 сотрудников", "name": "Технадзор/техсопровождение, 2–5 сотрудников", "unit": "чел.-месяц", "price_per_unit": 300000},
    {"code": "ОП-ЮРЗ 1 сотрудник", "name": "Юридическое сопровождение, 1 юрист", "unit": "чел.-месяц", "price_per_unit": 300000},
    {"code": "ОП-ЮРЗ 2-5 сотрудников", "name": "Юридическое сопровождение, 2–5 юристов", "unit": "чел.-месяц", "price_per_unit": 250000},
]

PRICE_LIST_STR = json.dumps(PRICE_LIST, ensure_ascii=False)

KP_SYSTEM_PROMPT = f"""Ты — эксперт по составлению коммерческих предложений для проектно-инжиниринговой компании.
На основе предоставленных документов (ТЗ, технические задания, файлы) составь детальное Коммерческое Предложение (КП) на русском языке.

ПРАЙС-ЛИСТ компании (используй ТОЛЬКО эти позиции для расчёта стоимости):
{PRICE_LIST_STR}

АЛГОРИТМ СОСТАВЛЕНИЯ КП (выполняй строго по шагам):

ШАГ 1 — Извлечь из ТЗ ключевые параметры:
- Тип объекта и наименование
- Вид работ: новое строительство (НС) / реконструкция (РЕК) / капремонт (КР)
- Площадь участка изысканий в га (найди в ТЗ явно или рассчитай из размеров)
- Строительный объём новых/реконструируемых сооружений в м³ (найди или оцени по аналогам)
- Наличие требований к экспертизам (ГЭЭ, госэкспертиза)
- Наличие авторского надзора, НТС

ШАГ 2 — Подобрать позиции из прайса:
ИНЖЕНЕРНЫЕ ИЗЫСКАНИЯ (для каждого вида изысканий):
- ИИ-ИГДИ (геодезия): объём = площадь_га, цена = 40000/га
- ИИ-ИЭИ (экология): объём = площадь_га, цена = 65000/га
- ИИ-ОБМ (обследование конструкций): если есть существующие сооружения, объём в м³, цена = 500/м³
- Инженерно-геологические, гидрометеорологические, археологические — если требуются в ТЗ, добавляй с ориентировочной ценой (нет в прайсе)

ПРОЕКТИРОВАНИЕ:
- Новое строительство ОКС:
  * до 100 м³ → "ОП-ОКС.НС до 100", 5000 руб/м³
  * 100–500 м³ → "ОП-ОКС.НС 100-500", 4000 руб/м³
  * 500–5000 м³ → "ОП-ОКС.НС 500-5000", 2000 руб/м³
  * более 5000 м³ → "ОП-ОКС.НС более 5000", 1500 руб/м³  ← для крупных объектов
- Реконструкция ОКС: "ОП-ОКС.РЕК ..." соответствующий диапазон

ЭКСПЕРТИЗЫ:
- СОПГЭ: 1 комплект = 200000 руб (всегда включай если требуется госэкспертиза)

ДОПОЛНИТЕЛЬНЫЕ (только если явно указано в ТЗ):
- ОП-АН (авторский надзор): объём = кол-во месяцев, 100000/мес
- ОП-НТС: объём = кол-во месяцев, 1500000/мес

ШАГ 3 — Рассчитать каждую строку:
total = volume × price_per_unit (ПРОВЕРИТЬ арифметику!)
Если объём из ТЗ не указан → ОЦЕНИ по аналогам (НЕ СТАВЬ 0!):
- Административное здание: ~5000-10000 м³
- Промышленный объект: ~10000-50000 м³
- Жилой дом: ~5000-15000 м³

ШАГ 4 — Посчитать total_sum:
Сложить ВСЕ строки total. Перепроверить сумму.

ПРИМЕР (для объекта 6.9 га, 25000 м³ НС):
- ИИ-ИГДИ: 6.9 × 40000 = 276000
- ИИ-ИЭИ: 6.9 × 65000 = 448500
- ОП-ОКС.НС более 5000: 25000 × 1500 = 37500000
- СОПГЭ: 1 × 200000 = 200000
- ОП-АН: 18 × 100000 = 1800000
- ИТОГО = 40224500

СТРУКТУРА КП:
1. Титульная страница: название, клиент, дата, объект, вид работ, основание, источник финансирования
2. Краткое резюме: суть проекта 3-5 предложений
3. Разделы работ: "Инженерные изыскания", "Проектирование", "Сопровождение экспертиз", "Дополнительные услуги"
4. В каждом разделе — таблица: Код | Наименование | Ед. | Объём | Цена за ед. | Итого
5. Итоговая стоимость с разбивкой: базовый пакет и опции
6. Сроки реализации по этапам
7. Условия оплаты (аванс 40%, промежуточные платежи, финальный расчёт)
8. Особые условия и исключения

Отвечай ТОЛЬКО JSON в формате:
{{
  "kp": {{
    "title": "...",
    "client": "...",
    "date": "...",
    "object_name": "...",
    "work_type": "...",
    "summary": "...",
    "sections": [{{"name": "...", "items": [{{"code": "...", "name": "...", "unit": "...", "volume": 0, "price_per_unit": 0, "total": 0, "notes": "..."}}]}}],
    "total_sum": 0,
    "total_sum_words": "...",
    "timeline": [{{"phase": "...", "duration": "...", "description": "..."}}],
    "payment_conditions": ["..."],
    "special_conditions": ["..."],
    "validity": "30 дней"
  }}
}}"""

ROADMAP_SYSTEM_PROMPT = """Ты — опытный менеджер проектов. На основе технического задания и коммерческого предложения составь подробную дорожную карту реализации проекта.

СТРУКТУРА ДОРОЖНОЙ КАРТЫ:
1. Обзор проекта: Цели, ключевые результаты
2. Этапы реализации: Детальная разбивка по фазам
3. Ключевые вехи: Контрольные точки и критерии приёмки
4. Риски и митигация: Основные риски и пути их снижения
5. Команда и ресурсы: Необходимые специалисты и ресурсы
6. Критический путь: Наиболее важные зависимости

Отвечай ТОЛЬКО JSON в формате:
{
  "roadmap": {
    "project_name": "...",
    "total_duration": "...",
    "overview": "...",
    "phases": [{"id": 1, "name": "...", "duration": "...", "start_week": 1, "end_week": 4, "tasks": ["..."], "deliverables": ["..."], "responsible": "..."}],
    "milestones": [{"week": 4, "name": "...", "criteria": "..."}],
    "risks": [{"risk": "...", "probability": "высокая/средняя/низкая", "impact": "высокий/средний/низкий", "mitigation": "..."}],
    "team": [{"role": "...", "count": 1, "tasks": "..."}],
    "critical_path": ["..."]
  }
}"""


def call_ai(system_prompt: str, user_message: str) -> str:
    api_key = os.environ.get('OPENROUTER_API_KEY', '')
    response = httpx.post(
        'https://openrouter.ai/api/v1/chat/completions',
        headers={
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json',
        },
        json={
            'model': 'deepseek/deepseek-chat',
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_message}
            ],
            'max_tokens': 8000,
            'temperature': 0.2,
        },
        timeout=180.0
    )
    result = response.json()
    if 'choices' not in result:
        error_msg = result.get('error', {}).get('message', str(result))
        raise Exception(f"OpenRouter error: {error_msg}")
    return result['choices'][0]['message']['content']


def extract_json(text: str) -> dict:
    text = text.strip()
    if '```json' in text:
        text = text.split('```json')[1].split('```')[0].strip()
    elif '```' in text:
        text = text.split('```')[1].split('```')[0].strip()
    return json.loads(text)


def process_job_async(job_id: str, action: str, combined_text: str, extra_prompt: str, kp_data: dict):
    """Фоновый поток: выполняет запрос к ИИ и сохраняет результат в БД"""
    conn = None
    try:
        if action == 'generate_kp':
            user_message = f"""СОДЕРЖИМОЕ ЗАГРУЖЕННЫХ ДОКУМЕНТОВ:
{combined_text}

{f'ДОПОЛНИТЕЛЬНЫЕ ТРЕБОВАНИЯ ОТ КЛИЕНТА: {extra_prompt}' if extra_prompt else ''}

Составь коммерческое предложение на основе этих материалов."""
            ai_response = call_ai(KP_SYSTEM_PROMPT, user_message)
        else:
            user_message = f"""СОДЕРЖИМОЕ ЗАГРУЖЕННЫХ ДОКУМЕНТОВ:
{combined_text}

{f'ДОПОЛНИТЕЛЬНЫЕ ТРЕБОВАНИЯ: {extra_prompt}' if extra_prompt else ''}
"""
            if kp_data:
                user_message += f"\nКОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ:\n{json.dumps(kp_data, ensure_ascii=False)}"
            user_message += "\n\nСоставь дорожную карту реализации проекта."
            ai_response = call_ai(ROADMAP_SYSTEM_PROMPT, user_message)

        result_data = extract_json(ai_response)
        result_json = json.dumps(result_data, ensure_ascii=False)

        conn = get_db()
        cur = conn.cursor()
        cur.execute(
            "UPDATE kp_jobs SET status='done', result=%s, updated_at=NOW() WHERE id=%s",
            (result_json, job_id)
        )
        conn.commit()
    except Exception as e:
        try:
            if conn is None:
                conn = get_db()
            cur = conn.cursor()
            cur.execute(
                "UPDATE kp_jobs SET status='error', error=%s, updated_at=NOW() WHERE id=%s",
                (str(e), job_id)
            )
            conn.commit()
        except Exception:
            pass
    finally:
        if conn:
            conn.close()


def handler(event: dict, context) -> dict:
    """Генерация КП через ИИ — асинхронная очередь (start_job → check_job) чтобы обойти таймаут"""

    if event.get('httpMethod') == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type, X-Authorization',
                'Access-Control-Max-Age': '86400'
            },
            'body': ''
        }

    if event.get('httpMethod') != 'POST':
        return {
            'statusCode': 405,
            'headers': {'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Method not allowed'})
        }

    body = json.loads(event.get('body', '{}'))
    action = body.get('action', '')

    CORS = {'Access-Control-Allow-Origin': '*', 'Content-Type': 'application/json'}

    if action == 'ping':
        return {'statusCode': 200, 'headers': CORS, 'body': json.dumps({'status': 'ok'})}

    # Проверка статуса задачи
    if action == 'check_job':
        job_id = body.get('job_id')
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT status, result, error FROM kp_jobs WHERE id=%s", (job_id,))
        row = cur.fetchone()
        conn.close()
        if not row:
            return {'statusCode': 404, 'headers': CORS, 'body': json.dumps({'error': 'Job not found'})}
        status, result, error = row
        resp = {'status': status}
        if status == 'done' and result:
            resp['data'] = json.loads(result)
        if status == 'error':
            resp['error'] = error
        return {'statusCode': 200, 'headers': CORS, 'body': json.dumps(resp, ensure_ascii=False)}

    # Запуск новой задачи (generate_kp или generate_roadmap)
    if action not in ('generate_kp', 'generate_roadmap'):
        return {'statusCode': 400, 'headers': CORS, 'body': json.dumps({'error': f'Unknown action: {action}'})}

    extra_prompt = body.get('extra_prompt', '')
    files_text = body.get('files_text', '')
    kp_data = body.get('kp_data', None)
    files_b64 = body.get('files_b64', [])

    # Парсинг файлов
    parsed_texts = []
    for f in files_b64:
        name = f.get('name', '')
        b64 = f.get('data', '')
        raw = base64.b64decode(b64)
        if name.lower().endswith('.pdf'):
            text = extract_text_from_pdf(raw)
        elif name.lower().endswith('.docx'):
            text = extract_text_from_docx(raw)
        else:
            text = raw.decode('utf-8', errors='replace')
        parsed_texts.append(f'=== ФАЙЛ: {name} ===\n{text}')

    combined_text = '\n\n'.join(parsed_texts) if parsed_texts else files_text

    # Сохранить задачу в БД
    input_data = json.dumps({'action': action, 'extra_prompt': extra_prompt}, ensure_ascii=False)
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO kp_jobs (action, status, input_data) VALUES (%s, 'processing', %s) RETURNING id",
        (action, input_data)
    )
    job_id = str(cur.fetchone()[0])
    conn.commit()
    conn.close()

    # Запустить фоновый поток
    t = threading.Thread(
        target=process_job_async,
        args=(job_id, action, combined_text, extra_prompt, kp_data),
        daemon=True
    )
    t.start()

    return {
        'statusCode': 202,
        'headers': CORS,
        'body': json.dumps({'job_id': job_id, 'status': 'processing'})
    }
